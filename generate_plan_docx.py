"""
generate_plan_docx.py
Converts PORTAL_PLAN.md into a formatted Word document.
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ── Brand colors ────────────────────────────────────────────────────────────
TEAL        = RGBColor(0x0F, 0x76, 0x6E)   # #0F766E  primary brand
TEAL_DARK   = RGBColor(0x08, 0x3D, 0x39)   # #083D39  dark accent
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_700    = RGBColor(0x37, 0x41, 0x51)
GRAY_400    = RGBColor(0x9C, 0xA3, 0xAF)
BLACK       = RGBColor(0x11, 0x18, 0x27)

MD_PATH   = "PORTAL_PLAN.md"
DOCX_PATH = "Vichakra_Portal_Plan.docx"


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex colour string (e.g. '0F766E')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_table_border(table):
    """Add thin borders to every cell of a table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "D1D5DB")
            tcBorders.append(border)
            tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, color=None, size=None, font="Outfit"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def heading_para(doc, text, level=1):
    """Add a styled heading that matches site branding."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(6)
    sizes   = {1: 22, 2: 16, 3: 13, 4: 11}
    colors  = {1: TEAL_DARK, 2: TEAL, 3: GRAY_700, 4: GRAY_700}
    add_run(para, text,
            bold=(level <= 2),
            color=colors.get(level, BLACK),
            size=sizes.get(level, 11))
    # Underline rule for H1
    if level == 1:
        para.paragraph_format.border_bottom = True
    return para


def body_para(doc, text, bullet=False, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    if bullet:
        para.style = doc.styles["List Bullet"]
    if indent:
        para.paragraph_format.left_indent = Inches(indent * 0.25)
    add_run(para, text, color=GRAY_700, size=10)
    return para


def code_para(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)   # blue-800
    return para


def add_divider(doc):
    para = doc.add_paragraph("─" * 90)
    run  = para.runs[0]
    run.font.color.rgb = GRAY_400
    run.font.size      = Pt(7)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)


def add_table_from_rows(doc, rows, header=True):
    """rows: list of lists of strings."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            is_header = header and r_idx == 0
            run = para.add_run(cell_text)
            run.font.name  = "Outfit"
            run.font.size  = Pt(9)
            run.bold       = is_header
            run.font.color.rgb = WHITE if is_header else GRAY_700
            if is_header:
                set_cell_bg(cell, "0F766E")
    doc.add_paragraph()   # spacing after table


# ── Cover page ───────────────────────────────────────────────────────────────

def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "VICHAKRA TECHNOLOGIES", bold=True, color=TEAL_DARK, size=28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "Admin Panel & Client Portal", bold=True, color=TEAL, size=18)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(desc, "Full Build Plan  ·  MERN Stack  ·  Phase-Step-Wise  ·  Dependency Flow",
            color=GRAY_400, size=11, italic=True)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "Version 1.0  ·  Author: Rohith  ·  Date: 2026-04-11",
            color=GRAY_400, size=10)

    doc.add_page_break()


# ── Markdown parser ──────────────────────────────────────────────────────────

def parse_md_to_docx(doc, md_text):
    lines       = md_text.splitlines()
    in_code     = False
    code_buf    = []
    table_rows  = []
    in_table    = False

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            # Remove separator rows (---|---)
            cleaned = [r for r in table_rows
                       if not all(c.strip().replace("-", "").replace(":", "") == "" for c in r)]
            if cleaned:
                add_table_from_rows(doc, cleaned)
        table_rows = []
        in_table   = False

    for line in lines:
        stripped = line.strip()

        # ── Code fence ──────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                for cl in code_buf:
                    code_para(doc, cl)
                code_buf = []
            else:
                in_code = True
                if in_table:
                    flush_table()
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Table rows ───────────────────────────────────────────────────────
        if stripped.startswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        # ── Skip horizontal rules ────────────────────────────────────────────
        if re.match(r"^---+$", stripped):
            add_divider(doc)
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.+)", stripped)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            # strip inline markdown from heading text
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`",        r"\1", text)
            heading_para(doc, text, level)
            continue

        # ── Bullet / checklist ───────────────────────────────────────────────
        m = re.match(r"^[-*]\s+\[[ xX]\]\s+(.+)", stripped)   # checklist
        if m:
            body_para(doc, "☐  " + m.group(1).strip(), bullet=False)
            continue

        m = re.match(r"^[-*]\s+(.+)", stripped)
        if m:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(1))
            text = re.sub(r"`(.+?)`",        r"\1", text)
            body_para(doc, text, bullet=True)
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.+)", stripped)
        if m:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(1))
            text = re.sub(r"`(.+?)`",        r"\1", text)
            body_para(doc, text, bullet=True)
            continue

        # ── Blank line ───────────────────────────────────────────────────────
        if stripped == "":
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        text = re.sub(r"`(.+?)`",        r"\1", text)
        body_para(doc, text)

    if in_table:
        flush_table()


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(MD_PATH):
        print(f"ERROR: '{MD_PATH}' not found. Run this script from the repo root.")
        return

    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Outfit"
    doc.styles["Normal"].font.size = Pt(10)

    add_cover(doc)
    parse_md_to_docx(doc, md_text)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
